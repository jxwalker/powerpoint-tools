import os
import logging
import asyncio
from typing import List, Tuple
from docx import Document
from fpdf import FPDF
import aiofiles

async def check_write_permissions(output_path: str):
    directory = os.path.dirname(output_path)
    if directory and not os.path.exists(directory):
        logging.error(f"Directory does not exist: {directory}")
        raise FileNotFoundError(f"Directory does not exist: {directory}")
    
    # Check write permissions for the directory or current working directory
    check_dir = directory if directory else '.'
    if not os.access(check_dir, os.W_OK):
        logging.error(f"No write permissions for directory: {check_dir}")
        raise PermissionError(f"No write permissions for directory: {check_dir}")

async def write_summary_to_word(summaries: List[Tuple[int, str, str]], output_path: str, summary_only: bool = False) -> None:
    await check_write_permissions(output_path)
    logging.info(f"Starting to write summary to Word document: {output_path}")
    logging.info(f"Number of summaries to write: {len(summaries)}")
    
    doc = Document()
    doc.add_heading("Presentation Notes", 0)
    
    for slide_number, summary, note in summaries:
        logging.debug(f"Writing slide {slide_number} to document")
        doc.add_heading(f"Slide {slide_number}", level=1)
        if summary:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(summary)
            logging.debug(f"Added summary for slide {slide_number}")
        if not summary_only and note:
            doc.add_heading("Notes", level=2)
            doc.add_paragraph(note)
            logging.debug(f"Added notes for slide {slide_number}")
    
    logging.info(f"Saving document to {output_path}")
    try:
        await asyncio.to_thread(doc.save, output_path)
        logging.info(f"Summary document saved to {output_path}")
    except Exception as e:
        logging.error(f"Failed to save Word document to {output_path}: {str(e)}")
        raise IOError(f"Failed to save Word document: {str(e)}") from e

async def write_summary_to_md(summaries: List[Tuple[int, str, str]], output_path: str, summary_only: bool = False) -> None:
    await check_write_permissions(output_path)
    try:
        async with aiofiles.open(output_path, "w") as file:
            await file.write("# Presentation Notes\n\n")
            for slide_number, summary, note in summaries:
                await file.write(f"## Slide {slide_number}\n\n")
                if summary:
                    await file.write("### Summary\n\n")
                    await file.write(summary + "\n\n")
                if not summary_only and note:
                    await file.write("### Notes\n\n")
                    await file.write(note + "\n\n")
        logging.info(f"Summary document saved to {output_path}")
    except IOError as e:
        logging.error(f"Failed to write Markdown file to {output_path}: {str(e)}")
        raise IOError(f"Failed to write Markdown file: {str(e)}") from e

async def write_summary_to_pdf(summaries: List[Tuple[int, str, str]], output_path: str, summary_only: bool = False) -> None:
    await check_write_permissions(output_path)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", style="B", size=16)
    pdf.cell(200, 10, txt="Presentation Notes", ln=True, align="C")
    for slide_number, summary, note in summaries:
        pdf.set_font("Arial", style="B", size=14)
        pdf.cell(200, 10, txt=f"Slide {slide_number}", ln=True, align="L")
        if summary:
            pdf.set_font("Arial", style="B", size=12)
            pdf.cell(200, 10, txt="Summary", ln=True, align="L")
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 5, txt=summary)
            pdf.ln(5)
        if not summary_only and note:
            pdf.set_font("Arial", style="B", size=12)
            pdf.cell(200, 10, txt="Notes", ln=True, align="L")
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 5, txt=note)
            pdf.ln(5)
        pdf.ln(10)  # Add some space between slides
    try:
        await asyncio.to_thread(pdf.output, output_path)
        logging.info(f"Summary document saved to {output_path}")
    except Exception as e:
        logging.error(f"Failed to save PDF to {output_path}: {str(e)}")
        raise IOError(f"Failed to save PDF: {str(e)}") from e

async def write_summary_to_txt(summaries: List[Tuple[int, str, str]], output_path: str, summary_only: bool = False) -> None:
    await check_write_permissions(output_path)
    try:
        async with aiofiles.open(output_path, "w") as file:
            await file.write("Presentation Notes\n\n")
            for slide_number, summary, note in summaries:
                await file.write(f"Slide {slide_number}\n")
                if summary:
                    await file.write("Summary:\n")
                    await file.write(summary + "\n\n")
                if not summary_only and note:
                    await file.write("Notes:\n")
                    await file.write(note + "\n\n")
        logging.info(f"Summary document saved to {output_path}")
    except IOError as e:
        logging.error(f"Failed to write text file to {output_path}: {str(e)}")
        raise IOError(f"Failed to write text file: {str(e)}") from e